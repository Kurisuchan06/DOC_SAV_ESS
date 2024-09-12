import tkinter as tk
from tkinter import ttk, messagebox
from docx import Document

# Fonction pour générer le document Word
def generer_document():
    # Récupérer les données des différents onglets
    nom = entry_nom.get()
    prenom = entry_prenom.get()
    adresse = entry_adresse.get()
    telephone = entry_telephone.get()
    email = entry_email.get()
    
    date_entree = entry_date_entree.get()
    date_fin = entry_date_fin.get()
    duree_totale = entry_duree_totale.get()
    rythme = entry_rythme.get()
    rythme_entreprise = entry_rythme_entreprise.get()
    volume_horaire = entry_horaire
    
    competence1 = var_competence1.get()
    competence2 = var_competence2.get()
    
    # Vérification des champs
    if not nom or not prenom or not date_entree or not date_fin or not rythme:
        messagebox.showwarning("Erreur", "Veuillez remplir tous les champs.")
        return
    
    # Créer un document Word
    doc = Document()
    doc.add_heading('Informations sur la Formation', 0)

    # Ajouter les informations personnelles
    doc.add_paragraph(f"Nom : {nom}")
    doc.add_paragraph(f"Prénom : {prenom}")
    doc.add_paragraph(f"Adresse : {adresse}")
    doc.add_paragraph(f"Téléphone : {telephone}")
    doc.add_paragraph(f"Email : {email}")

    # Ajouter les informations sur la formation
    doc.add_paragraph(f"Date d'entrée : {date_entree}")
    doc.add_paragraph(f"Date de fin : {date_fin}")
    doc.add_paragraph(f"Durée totale : {duree_totale}")
    doc.add_paragraph(f"Rythme hebdomadaire : {rythme}")
    doc.add_paragraph(f"Rythme en entreprise : {rythme_entreprise}")
    doc.add_paragraph(f"Volume horaire : {duree_totale}")


    # Ajouter les compétences cochées
    doc.add_heading('Compétences', level=1)
    if competence1:
        doc.add_paragraph("Compétence 1 : Apprise")
    if competence2:
        doc.add_paragraph("Compétence 2 : Apprise")

    # Sauvegarder le document
    doc.save('formulaire_formation.docx')
    messagebox.showinfo("Succès", "Le document a été généré avec succès.")

# Création de la fenêtre principale
root = tk.Tk()
root.title("Générateur de Document Word - Formation")

# Création du Notebook pour les onglets
notebook = ttk.Notebook(root)
notebook.grid(row=0, column=0, padx=10, pady=10)

# ------------- Onglet 1 : Informations personnelles --------------
frame_infos_personnelles = ttk.Frame(notebook)
notebook.add(frame_infos_personnelles, text="Informations personnelles")

# Champs pour les informations personnelles
tk.Label(frame_infos_personnelles, text="Nom :").grid(row=0, column=0)
entry_nom = tk.Entry(frame_infos_personnelles)
entry_nom.grid(row=0, column=1)

tk.Label(frame_infos_personnelles, text="Prénom :").grid(row=1, column=0)
entry_prenom = tk.Entry(frame_infos_personnelles)
entry_prenom.grid(row=1, column=1)

tk.Label(frame_infos_personnelles, text="Adresse :").grid(row=2, column=0)
entry_adresse = tk.Entry(frame_infos_personnelles)
entry_adresse.grid(row=2, column=1)

tk.Label(frame_infos_personnelles, text="Numéro de téléphone :").grid(row=3, column=0)
entry_telephone = tk.Entry(frame_infos_personnelles)
entry_telephone.grid(row=3, column=1)

tk.Label(frame_infos_personnelles, text="Email :").grid(row=4, column=0)
entry_email = tk.Entry(frame_infos_personnelles)
entry_email.grid(row=4, column=1)

# ------------- Onglet 2 : Compétences disponibles --------------
frame_competences = ttk.Frame(notebook)
notebook.add(frame_competences, text="Compétences")

# Cases à cocher pour les compétences
var_competence1 = tk.BooleanVar()
var_competence2 = tk.BooleanVar()

tk.Checkbutton(frame_competences, text="Compétence 1", variable=var_competence1).grid(row=0, column=0)
tk.Checkbutton(frame_competences, text="Compétence 2", variable=var_competence2).grid(row=1, column=0)

# ------------- Onglet 3 : Informations sur la formation --------------
frame_infos_formation = ttk.Frame(notebook)
notebook.add(frame_infos_formation, text="Informations Formation")

# Champs pour les informations de la formation
tk.Label(frame_infos_formation, text="Date d'entrée :").grid(row=0, column=0)
entry_date_entree = tk.Entry(frame_infos_formation)
entry_date_entree.grid(row=0, column=1)

tk.Label(frame_infos_formation, text="Date de fin :").grid(row=1, column=0)
entry_date_fin = tk.Entry(frame_infos_formation)
entry_date_fin.grid(row=1, column=1)

tk.Label(frame_infos_formation, text="Rythme hebdomadaire :").grid(row=2, column=0)
entry_rythme = tk.Entry(frame_infos_formation)
entry_rythme.grid(row=2, column=1)

# Bouton pour générer le document
btn_generer = tk.Button(root, text="Générer le document", command=generer_document)
btn_generer.grid(row=1, column=0, padx=10, pady=10)

# Lancer l'application
root.mainloop()